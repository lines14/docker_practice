# document 1:
# petition for familiarization with the case materials

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from modules import template

async def data_print(state):
    async with state.proxy() as data:
        values = list(data.values())
        for val in values:
            print(val)

    def make_table_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def make_table_columns_align_right(*columns):
        for column in columns:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def set_table_rows_height_0_5(*rows):
        for row in rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.5)

    # font = paragraph.runs[0].font
    # font.size= Pt(10)
    # style = document.styles['Normal']
    # style.paragraph_format.line_spacing = Pt(8)

    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2) #Верхний отступ
    section.bottom_margin = Cm(2) #Нижний отступ
    section.left_margin = Cm(3) #Отступ слева
    section.right_margin = Cm(2) #Отступ справа

    records = (
        (template.VAR1, values[1]),
        (template.VAR3, values[2]),
        (template.VAR5, values[3]),
        (template.VAR7, values[4]),
        (template.VAR9, values[5]),
        (template.VAR11, values[6]),
        (template.VAR13, values[7]),
        (template.VAR15, values[8]),
        (template.VAR17, template.VAR18),
        (template.VAR19, values[9])
    )

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = template.VAR21
    hdr_cells[1].text = values[0]
    for first_col, second_col in records:
        row_cells = table.add_row().cells
        row_cells[0].text = first_col
        row_cells[1].text = second_col

    make_table_rows_bold(table.rows[0], table.rows[2], table.rows[7])
    make_table_columns_align_right(table.columns[0])
    set_table_rows_height_0_5(table.rows[9])

    para = document.add_paragraph(values[10])
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_after = Pt(8)

    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    para.add_run(template.VAR24).bold = True


    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    para.add_run(template.VAR25).bold = True


    para = document.add_paragraph(f'\t{values[11]}')
    para.paragraph_format.space_after = Pt(8)

    # para = document.add_paragraph(f'\t{template.VAR27}')
    # para.paragraph_format.space_after = Pt(8)

    # para = document.add_paragraph(f'\t{template.VAR28}')
    # para.paragraph_format.space_after = Pt(8)

    # para = document.add_paragraph(template.VAR29)
    # para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # para.paragraph_format.space_after = Pt(8)

    # para = document.add_paragraph(f'\t{template.VAR30}\n')
    # para.paragraph_format.space_after = Pt(8)

    records = (
    (values[12], values[13]),
    )

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = template.VAR31
    hdr_cells[1].text = template.VAR32
    for first_col, second_col in records:
        row_cells = table.add_row().cells
        row_cells[0].text = first_col
        row_cells[1].text = second_col

    make_table_rows_bold(table.rows[1])
    make_table_columns_align_right(table.columns[1])

    document.save('/home/lines14/projects/judicial_telegram_bot/documents/judicial_writer_1.docx')